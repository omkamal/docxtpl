"""Script to create the resume_template.docx file."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_resume_template():
    """Create a professional resume template with Jinja2 placeholders."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header - Name
    header = doc.add_heading(level=0)
    run = header.add_run("{{ full_name }}")
    run.font.size = Pt(24)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("{{ email }} | {{ phone }} | {{ location }}")

    doc.add_paragraph()  # Spacer

    # Professional Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("{{ summary }}")

    # Experience Section
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("{%- for job in experience %}")

    exp_title = doc.add_paragraph()
    exp_title.add_run("{{ job.title }}").bold = True
    exp_title.add_run(" at ")
    exp_title.add_run("{{ job.company }}").bold = True

    exp_dates = doc.add_paragraph()
    exp_dates.add_run("{{ job.start_date }} - {{ job.end_date }}")

    doc.add_paragraph("{{ job.description }}")
    doc.add_paragraph("{%- endfor %}")

    # Education Section
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{%- for edu in education %}")

    edu_info = doc.add_paragraph()
    edu_info.add_run("{{ edu.degree }}").bold = True
    edu_info.add_run(" - {{ edu.school }}, {{ edu.year }}")

    doc.add_paragraph("{%- endfor %}")

    # Skills Section
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("{{ skills }}")

    # Optional Certifications Section
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("{%- for cert in certifications %}")
    doc.add_paragraph("{{ cert.name }} - {{ cert.issuer }} ({{ cert.year }})")
    doc.add_paragraph("{%- endfor %}")

    # Save the template
    doc.save("resume_template.docx")
    print("Created resume_template.docx")


if __name__ == "__main__":
    create_resume_template()
