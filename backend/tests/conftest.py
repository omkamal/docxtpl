"""Test fixtures and configuration."""

import os
import tempfile
import pytest
from docx import Document
from httpx import AsyncClient, ASGITransport

from app.main import app


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def simple_template(temp_dir):
    """Create a simple template with basic fields."""
    doc = Document()
    doc.add_paragraph("Name: {{ name }}")
    doc.add_paragraph("Email: {{ email }}")
    doc.add_paragraph("Company: {{ company }}")

    path = os.path.join(temp_dir, "simple_template.docx")
    doc.save(path)
    return path


@pytest.fixture
def resume_template(temp_dir):
    """Create a resume template with fields and loops."""
    doc = Document()

    # Header
    doc.add_heading("{{ full_name }}", 0)
    doc.add_paragraph("{{ email }} | {{ phone }} | {{ location }}")

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ summary }}")

    # Experience section with loop
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("{%- for job in experience %}")
    doc.add_paragraph("{{ job.title }} at {{ job.company }}")
    doc.add_paragraph("{{ job.start_date }} - {{ job.end_date }}")
    doc.add_paragraph("{{ job.description }}")
    doc.add_paragraph("{%- endfor %}")

    # Education section with loop
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{%- for edu in education %}")
    doc.add_paragraph("{{ edu.degree }} - {{ edu.school }}")
    doc.add_paragraph("{{ edu.year }}")
    doc.add_paragraph("{%- endfor %}")

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("{{ skills }}")

    path = os.path.join(temp_dir, "resume_template.docx")
    doc.save(path)
    return path


@pytest.fixture
def template_with_table(temp_dir):
    """Create a template with table row loops."""
    doc = Document()
    doc.add_paragraph("Invoice for {{ customer_name }}")

    # Add a table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Quantity"
    header_cells[2].text = "Price"

    # Template row
    template_row = table.rows[1].cells
    template_row[0].text = "{%tr for item in items %}{{ item.name }}"
    template_row[1].text = "{{ item.quantity }}"
    template_row[2].text = "{{ item.price }}{%tr endfor %}"

    path = os.path.join(temp_dir, "table_template.docx")
    doc.save(path)
    return path


@pytest.fixture
def single_brace_template(temp_dir):
    """Create a template with single-brace {field} syntax."""
    doc = Document()
    doc.add_paragraph("Name: {name}")
    doc.add_paragraph("Email: {email}")
    doc.add_paragraph("Company: {company}")

    path = os.path.join(temp_dir, "single_brace_template.docx")
    doc.save(path)
    return path


@pytest.fixture
def template_with_sections(temp_dir):
    """Create a template with conditional and list sections."""
    doc = Document()
    doc.add_paragraph("Name: {name}")
    doc.add_paragraph("Email: {email}")

    # Conditional section (no nested fields)
    doc.add_paragraph("{#disclaimer}This is a legal disclaimer.{/disclaimer}")

    # List section with nested fields
    doc.add_paragraph("{#products}")
    doc.add_paragraph("Product: {product_name} - Price: {price}")
    doc.add_paragraph("{/products}")

    path = os.path.join(temp_dir, "sections_template.docx")
    doc.save(path)
    return path


@pytest.fixture
async def client():
    """Create an async test client."""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac
