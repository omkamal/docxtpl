"""Tests specifically for the resume_template.docx file."""

import os
import io
import pytest
from docx import Document

from app.template_parser import parse_docx_template
from app.document_generator import generate_document


# Path to the actual resume template
RESUME_TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
    "resume_template.docx"
)


class TestResumeTemplateFile:
    """Tests for the actual resume_template.docx file."""

    @pytest.fixture
    def resume_template_path(self):
        """Return path to the actual resume template."""
        if not os.path.exists(RESUME_TEMPLATE_PATH):
            pytest.skip("resume_template.docx not found")
        return RESUME_TEMPLATE_PATH

    def test_template_exists(self, resume_template_path):
        """Test that the resume template file exists."""
        assert os.path.exists(resume_template_path)

    def test_parse_resume_template(self, resume_template_path):
        """Test parsing the actual resume template."""
        result = parse_docx_template(resume_template_path)

        # Check expected fields
        assert "full_name" in result["fields"]
        assert "email" in result["fields"]
        assert "phone" in result["fields"]
        assert "location" in result["fields"]
        assert "summary" in result["fields"]
        assert "skills" in result["fields"]

        # Check expected loops
        assert "experience" in result["loops"]
        assert "education" in result["loops"]
        assert "certifications" in result["loops"]

        # Check experience loop fields
        exp_fields = result["loops"]["experience"]
        assert "title" in exp_fields
        assert "company" in exp_fields
        assert "start_date" in exp_fields
        assert "end_date" in exp_fields
        assert "description" in exp_fields

        # Check education loop fields
        edu_fields = result["loops"]["education"]
        assert "degree" in edu_fields
        assert "school" in edu_fields
        assert "year" in edu_fields

        # Check certifications loop fields
        cert_fields = result["loops"]["certifications"]
        assert "name" in cert_fields
        assert "issuer" in cert_fields
        assert "year" in cert_fields

    def test_generate_full_resume(self, resume_template_path):
        """Test generating a complete resume."""
        data = {
            "full_name": "Sarah Johnson",
            "email": "sarah.johnson@email.com",
            "phone": "(555) 123-4567",
            "location": "San Francisco, CA",
            "summary": "Seasoned software engineer with 10+ years of experience building scalable web applications and leading cross-functional teams.",
            "skills": "Python, JavaScript, TypeScript, React, Node.js, PostgreSQL, AWS, Docker, Kubernetes",
            "experience": [
                {
                    "title": "Senior Software Engineer",
                    "company": "TechCorp Inc.",
                    "start_date": "January 2020",
                    "end_date": "Present",
                    "description": "Led development of microservices architecture serving 1M+ daily users. Mentored junior developers and established coding standards."
                },
                {
                    "title": "Software Engineer",
                    "company": "StartupXYZ",
                    "start_date": "March 2017",
                    "end_date": "December 2019",
                    "description": "Built real-time data processing pipeline handling 100K events/second. Reduced infrastructure costs by 40%."
                },
                {
                    "title": "Junior Developer",
                    "company": "WebAgency",
                    "start_date": "June 2014",
                    "end_date": "February 2017",
                    "description": "Developed responsive web applications for enterprise clients using React and Node.js."
                }
            ],
            "education": [
                {
                    "degree": "M.S. Computer Science",
                    "school": "Stanford University",
                    "year": "2014"
                },
                {
                    "degree": "B.S. Computer Science",
                    "school": "UC Berkeley",
                    "year": "2012"
                }
            ],
            "certifications": [
                {
                    "name": "AWS Solutions Architect Professional",
                    "issuer": "Amazon Web Services",
                    "year": "2022"
                },
                {
                    "name": "Certified Kubernetes Administrator",
                    "issuer": "CNCF",
                    "year": "2021"
                }
            ]
        }

        result = generate_document(resume_template_path, data)

        # Verify the document was generated
        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        # Check all fields were filled
        assert "Sarah Johnson" in text
        assert "sarah.johnson@email.com" in text
        assert "(555) 123-4567" in text
        assert "San Francisco, CA" in text
        assert "Seasoned software engineer" in text
        assert "Python, JavaScript" in text

        # Check experience entries
        assert "Senior Software Engineer" in text
        assert "TechCorp Inc." in text
        assert "Software Engineer" in text
        assert "StartupXYZ" in text
        assert "Junior Developer" in text
        assert "WebAgency" in text

        # Check education entries
        assert "M.S. Computer Science" in text
        assert "Stanford University" in text
        assert "B.S. Computer Science" in text
        assert "UC Berkeley" in text

        # Check certifications
        assert "AWS Solutions Architect Professional" in text
        assert "Amazon Web Services" in text
        assert "Certified Kubernetes Administrator" in text
        assert "CNCF" in text

        # Verify no placeholders remain
        assert "{{" not in text
        assert "}}" not in text
        assert "{%" not in text
        assert "%}" not in text

    def test_generate_minimal_resume(self, resume_template_path):
        """Test generating a resume with minimal data."""
        data = {
            "full_name": "John Minimal",
            "email": "john@minimal.com",
            "experience": [
                {
                    "title": "Developer",
                    "company": "Company",
                    "start_date": "2023",
                    "end_date": "Present",
                    "description": "Development work"
                }
            ]
        }

        result = generate_document(resume_template_path, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "John Minimal" in text
        assert "john@minimal.com" in text
        assert "Developer" in text

        # Unfilled fields should be cleaned
        assert "{{" not in text
        assert "}}" not in text

    def test_generate_resume_empty_sections(self, resume_template_path):
        """Test generating a resume with some empty sections."""
        data = {
            "full_name": "Empty Sections Test",
            "email": "test@test.com",
            "phone": "555-0000",
            "location": "Test City",
            "summary": "Test summary",
            "skills": "Test skills",
            "experience": [],  # Empty
            "education": [],   # Empty
            "certifications": []  # Empty
        }

        result = generate_document(resume_template_path, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Empty Sections Test" in text
        assert "test@test.com" in text

        # Loop placeholders should be cleaned
        assert "{%" not in text
        assert "%}" not in text
