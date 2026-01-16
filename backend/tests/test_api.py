"""Tests for API endpoints."""

import io
import pytest
from docx import Document


class TestHealthEndpoint:
    """Tests for health check endpoint."""

    async def test_health_check(self, client):
        """Test health endpoint returns healthy status."""
        response = await client.get("/api/health")

        assert response.status_code == 200
        assert response.json() == {"status": "healthy"}


class TestUploadEndpoint:
    """Tests for template upload endpoint."""

    async def test_upload_simple_template(self, client, simple_template):
        """Test uploading a simple template."""
        with open(simple_template, "rb") as f:
            response = await client.post(
                "/api/upload",
                files={"file": ("template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        assert response.status_code == 200
        data = response.json()

        assert "template_id" in data
        assert "fields" in data
        assert "loops" in data
        assert "name" in data["fields"]
        assert "email" in data["fields"]
        assert "company" in data["fields"]

    async def test_upload_resume_template(self, client, resume_template):
        """Test uploading a resume template with loops."""
        with open(resume_template, "rb") as f:
            response = await client.post(
                "/api/upload",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        assert response.status_code == 200
        data = response.json()

        assert "full_name" in data["fields"]
        assert "experience" in data["loops"]
        assert "education" in data["loops"]

    async def test_upload_invalid_file_type(self, client, temp_dir):
        """Test uploading a non-docx file."""
        import os
        txt_file = os.path.join(temp_dir, "test.txt")
        with open(txt_file, "w") as f:
            f.write("Not a docx file")

        with open(txt_file, "rb") as f:
            response = await client.post(
                "/api/upload",
                files={"file": ("test.txt", f, "text/plain")}
            )

        assert response.status_code == 400
        assert "docx" in response.json()["detail"].lower()

    async def test_upload_no_file(self, client):
        """Test uploading without a file."""
        response = await client.post("/api/upload")

        assert response.status_code == 422  # Validation error


class TestGenerateEndpoint:
    """Tests for document generation endpoint."""

    async def test_generate_simple_document(self, client, simple_template):
        """Test generating a simple document."""
        # First upload the template
        with open(simple_template, "rb") as f:
            upload_response = await client.post(
                "/api/upload",
                files={"file": ("template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        template_id = upload_response.json()["template_id"]

        # Generate document
        response = await client.post(
            "/api/generate",
            json={
                "template_id": template_id,
                "data": {
                    "name": "Test User",
                    "email": "test@example.com",
                    "company": "Test Corp"
                }
            }
        )

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Verify content
        doc = Document(io.BytesIO(response.content))
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Test User" in text
        assert "test@example.com" in text
        assert "Test Corp" in text

    async def test_generate_resume(self, client, resume_template):
        """Test generating a resume with loops."""
        # Upload template
        with open(resume_template, "rb") as f:
            upload_response = await client.post(
                "/api/upload",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        template_id = upload_response.json()["template_id"]

        # Generate document
        response = await client.post(
            "/api/generate",
            json={
                "template_id": template_id,
                "data": {
                    "full_name": "Jane Developer",
                    "email": "jane@dev.com",
                    "phone": "555-0123",
                    "location": "San Francisco, CA",
                    "summary": "Full-stack developer",
                    "skills": "Python, React, Docker",
                    "experience": [
                        {
                            "title": "Lead Developer",
                            "company": "BigTech",
                            "start_date": "2021",
                            "end_date": "Present",
                            "description": "Leading development"
                        }
                    ],
                    "education": [
                        {
                            "degree": "MS CS",
                            "school": "Tech University",
                            "year": "2021"
                        }
                    ]
                }
            }
        )

        assert response.status_code == 200

        doc = Document(io.BytesIO(response.content))
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Jane Developer" in text
        assert "Lead Developer" in text
        assert "BigTech" in text
        assert "MS CS" in text
        assert "Tech University" in text

    async def test_generate_with_partial_data(self, client, simple_template):
        """Test generating with partial data - unfilled fields should be cleaned."""
        with open(simple_template, "rb") as f:
            upload_response = await client.post(
                "/api/upload",
                files={"file": ("template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        template_id = upload_response.json()["template_id"]

        response = await client.post(
            "/api/generate",
            json={
                "template_id": template_id,
                "data": {
                    "name": "Only Name Provided"
                    # email and company not provided
                }
            }
        )

        assert response.status_code == 200

        doc = Document(io.BytesIO(response.content))
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Only Name Provided" in text
        # Unfilled placeholders should be removed
        assert "{{" not in text
        assert "}}" not in text

    async def test_generate_invalid_template_id(self, client):
        """Test generating with invalid template ID."""
        response = await client.post(
            "/api/generate",
            json={
                "template_id": "nonexistent-id",
                "data": {"name": "Test"}
            }
        )

        assert response.status_code == 404
        assert "not found" in response.json()["detail"].lower()

    async def test_generate_empty_loops(self, client, resume_template):
        """Test generating with empty loop data."""
        with open(resume_template, "rb") as f:
            upload_response = await client.post(
                "/api/upload",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )

        template_id = upload_response.json()["template_id"]

        response = await client.post(
            "/api/generate",
            json={
                "template_id": template_id,
                "data": {
                    "full_name": "Empty Loops Test",
                    "experience": [],  # Empty loop
                    "education": []    # Empty loop
                }
            }
        )

        assert response.status_code == 200

        doc = Document(io.BytesIO(response.content))
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Empty Loops Test" in text
        # Loop placeholders should be cleaned
        assert "{%" not in text
        assert "%}" not in text
