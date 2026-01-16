"""Tests for document_generator module."""

import pytest
from docx import Document

from app.document_generator import generate_document, clean_unused_placeholders


class TestGenerateDocument:
    """Tests for generate_document function."""

    def test_simple_generation(self, simple_template):
        """Test generating a document with simple fields."""
        data = {
            "name": "John Doe",
            "email": "john@example.com",
            "company": "Acme Inc"
        }

        result = generate_document(simple_template, data)

        # Read the generated document
        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "John Doe" in text
        assert "john@example.com" in text
        assert "Acme Inc" in text

    def test_partial_data(self, simple_template):
        """Test generating with only some fields filled."""
        data = {
            "name": "Jane Doe",
            # email and company are omitted
        }

        result = generate_document(simple_template, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Jane Doe" in text
        # Placeholders should be cleaned
        assert "{{" not in text
        assert "}}" not in text

    def test_empty_data(self, simple_template):
        """Test generating with no data."""
        data = {}

        result = generate_document(simple_template, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        # All placeholders should be cleaned
        assert "{{" not in text
        assert "}}" not in text

    def test_resume_with_loops(self, resume_template):
        """Test generating a resume with loop data."""
        data = {
            "full_name": "Alice Smith",
            "email": "alice@example.com",
            "phone": "555-1234",
            "location": "New York, NY",
            "summary": "Experienced software developer",
            "skills": "Python, JavaScript, React",
            "experience": [
                {
                    "title": "Senior Developer",
                    "company": "Tech Corp",
                    "start_date": "2020",
                    "end_date": "Present",
                    "description": "Led development team"
                },
                {
                    "title": "Developer",
                    "company": "Startup Inc",
                    "start_date": "2018",
                    "end_date": "2020",
                    "description": "Built web applications"
                }
            ],
            "education": [
                {
                    "degree": "BS Computer Science",
                    "school": "State University",
                    "year": "2018"
                }
            ]
        }

        result = generate_document(resume_template, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Alice Smith" in text
        assert "alice@example.com" in text
        assert "Senior Developer" in text
        assert "Tech Corp" in text
        assert "Developer" in text
        assert "Startup Inc" in text
        assert "BS Computer Science" in text
        assert "State University" in text

    def test_empty_loop_items_filtered(self, resume_template):
        """Test that empty loop items are filtered out."""
        data = {
            "full_name": "Bob Jones",
            "experience": [
                {"title": "Developer", "company": "Corp", "start_date": "2020", "end_date": "Present", "description": "Work"},
                {"title": "", "company": "", "start_date": "", "end_date": "", "description": ""},  # Empty - should be filtered
            ],
            "education": []  # Empty list
        }

        result = generate_document(resume_template, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Bob Jones" in text
        assert "Developer" in text

    def test_special_characters(self, simple_template):
        """Test generating with special characters (unicode)."""
        data = {
            "name": "José García",
            "email": "josé@ejemplo.com",
            "company": "Café Company"
        }

        result = generate_document(simple_template, data)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "José García" in text
        assert "josé@ejemplo.com" in text
        assert "Café Company" in text


class TestCleanUnusedPlaceholders:
    """Tests for clean_unused_placeholders function."""

    def test_cleans_simple_placeholders(self, temp_dir):
        """Test cleaning of simple placeholders."""
        import os
        from docx import Document
        import io

        # Create a document with placeholders
        doc = Document()
        doc.add_paragraph("Hello {{ name }}!")
        doc.add_paragraph("{{ unfilled_field }}")

        input_path = os.path.join(temp_dir, "test.docx")
        doc.save(input_path)

        output = io.BytesIO()
        clean_unused_placeholders(input_path, output)
        output.seek(0)

        cleaned_doc = Document(output)
        text = "\n".join([p.text for p in cleaned_doc.paragraphs])

        assert "{{" not in text
        assert "}}" not in text

    def test_cleans_control_statements(self, temp_dir):
        """Test cleaning of Jinja2 control statements."""
        import os
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("{% for item in items %}")
        doc.add_paragraph("{{ item.name }}")
        doc.add_paragraph("{% endfor %}")

        input_path = os.path.join(temp_dir, "test.docx")
        doc.save(input_path)

        output = io.BytesIO()
        clean_unused_placeholders(input_path, output)
        output.seek(0)

        cleaned_doc = Document(output)
        text = "\n".join([p.text for p in cleaned_doc.paragraphs])

        assert "{%" not in text
        assert "%}" not in text
        assert "{{" not in text
        assert "}}" not in text
