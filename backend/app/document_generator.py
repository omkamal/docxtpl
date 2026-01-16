"""Document generator using docxtpl."""

import io
import re
import zipfile
import shutil
from docxtpl import DocxTemplate
from docx import Document


def clean_unused_placeholders(doc_path: str, output: io.BytesIO) -> None:
    """Remove any remaining Jinja2 placeholders from the generated document.

    This handles cases where optional fields weren't provided.
    """
    doc = Document(doc_path)

    # Pattern to match Jinja2 placeholders
    placeholder_pattern = r'\{\{[^}]+\}\}|\{%[^%]+%\}'

    def clean_text(text: str) -> str:
        return re.sub(placeholder_pattern, '', text)

    # Clean paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = clean_text(run.text)

    # Clean tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = clean_text(run.text)

    # Clean headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = clean_text(run.text)
        if section.footer:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = clean_text(run.text)

    doc.save(output)


def _process_sections(content: str, data: dict) -> str:
    """Process {#section}...{/section} patterns in content.

    Handles:
    - Conditionals: Keep content if data[section] is truthy, else remove
    - Lists: Repeat content for each item in data[section] array
    """
    # Process sections (must be done before simple field replacement)
    section_pattern = r'\{#(\w+)\}(.*?)\{/\1\}'

    def replace_section(match):
        section_name = match.group(1)
        section_content = match.group(2)
        section_data = data.get(section_name)

        if section_data is None or section_data is False:
            # Remove section entirely
            return ''
        elif isinstance(section_data, list):
            # List section: repeat for each item
            result_parts = []
            for item in section_data:
                if isinstance(item, dict):
                    item_content = section_content
                    # Replace nested fields with item values
                    for field, value in item.items():
                        if value:
                            pattern = r'(?<!\{)\{' + re.escape(field) + r'\}(?!\})'
                            item_content = re.sub(pattern, str(value), item_content)
                    result_parts.append(item_content)
            return ''.join(result_parts)
        elif section_data is True or section_data:
            # Conditional section with truthy value: keep content, remove tags
            return section_content

        return ''

    return re.sub(section_pattern, replace_section, content, flags=re.DOTALL)


def _generate_single_brace(template_path: str, data: dict) -> io.BytesIO:
    """Generate document using simple string replacement for {field} syntax.

    Handles:
    - Simple fields: {field} replaced with data[field]
    - Conditionals: {#section}...{/section} shown if data[section] is truthy
    - Lists: {#section}...{/section} repeated for each item in data[section]

    Args:
        template_path: Path to the docx template file
        data: Dictionary of field values to fill in

    Returns:
        BytesIO containing the generated document
    """
    output = io.BytesIO()

    # Read the template as a ZIP file
    with zipfile.ZipFile(template_path, 'r') as template_zip:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as output_zip:
            for item in template_zip.namelist():
                content = template_zip.read(item)

                # Only process XML files
                if item.endswith('.xml'):
                    content_str = content.decode('utf-8')

                    # Process sections first (conditionals and lists)
                    content_str = _process_sections(content_str, data)

                    # Replace simple {field} patterns with data values
                    for field, value in data.items():
                        # Skip list/dict values (handled by sections)
                        if isinstance(value, (list, dict)):
                            continue
                        if value:  # Only replace non-empty values
                            pattern = r'(?<!\{)\{' + re.escape(field) + r'\}(?!\})'
                            content_str = re.sub(pattern, str(value), content_str)

                    # Remove any remaining unfilled {field} placeholders
                    content_str = re.sub(
                        r'(?<!\{)\{[a-zA-Z_][a-zA-Z0-9_]*\}(?!\})',
                        '',
                        content_str
                    )

                    # Remove any remaining section tags
                    content_str = re.sub(r'\{[#/]\w+\}', '', content_str)

                    content = content_str.encode('utf-8')

                output_zip.writestr(item, content)

    output.seek(0)
    return output


def generate_document(template_path: str, data: dict, syntax: str = 'jinja2') -> io.BytesIO:
    """Generate a filled document from a template and data.

    Args:
        template_path: Path to the docx template file
        data: Dictionary of field values to fill in
        syntax: Template syntax type ('jinja2' or 'single_brace')

    Returns:
        BytesIO containing the generated document
    """
    # Handle single-brace syntax separately
    if syntax == 'single_brace':
        return _generate_single_brace(template_path, data)

    # Jinja2 syntax - use docxtpl
    # Load the template
    doc = DocxTemplate(template_path)

    # Filter out empty values and empty lists
    filtered_data = {}
    for key, value in data.items():
        if isinstance(value, list):
            # Filter out items with all empty values
            filtered_list = []
            for item in value:
                if isinstance(item, dict):
                    # Keep items that have at least one non-empty value
                    if any(v for v in item.values() if v):
                        filtered_list.append(item)
                elif item:
                    filtered_list.append(item)
            if filtered_list:
                filtered_data[key] = filtered_list
        elif value:  # Only include non-empty values
            filtered_data[key] = value

    # Render the template with the data
    doc.render(filtered_data)

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    # Clean any remaining placeholders
    temp_output = io.BytesIO()
    clean_unused_placeholders(output, temp_output)
    temp_output.seek(0)

    return temp_output
